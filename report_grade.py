"""쇠퇴 등급 + 상권 + 빈점포 통합 리포트(docx). 등급은 원자료 그대로 제시(방향 단정 안 함)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
NAVY=RGBColor(0x1F,0x3A,0x5F); GRAY=RGBColor(0x55,0x55,0x55); RED=RGBColor(0xC0,0x39,0x2B)

def build(addr, geo, comm, grades, path, vacancy=None):
    doc=Document()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("대상지 쇠퇴 진단 리포트"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
    s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr=s.add_run(addr); sr.font.size=Pt(12); sr.font.color.rgb=GRAY
    m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
    mr=m.add_run(f"좌표 {geo['lat']:.5f}, {geo['lon']:.5f}  ·  생성일 {date.today()}  ·  데이터: 국토부 쇠퇴진단 등급 + 소상공인 상가정보 (모두 실데이터)")
    mr.font.size=Pt(9); mr.font.color.rgb=GRAY
    doc.add_paragraph()
    # 1. 상권
    _h(doc,"1. 상권 현황 (반경 500m · 실데이터)")
    doc.add_paragraph(f"영업 점포 수: {comm['total_stores']}개", style="List Bullet")
    top=", ".join(f"{k} {v}" for k,v in list(comm["by_major_category"].items())[:8])
    doc.add_paragraph(f"업종 구성: {top}", style="List Bullet")
    doc.add_paragraph(f"참고: {comm['note']}", style="List Bullet")
    doc.add_paragraph()
    # 2. 빈 점포/공실
    _h(doc,"2. 빈 점포 · 공실 (시점비교 방식)")
    if vacancy and vacancy.get("has_prev"):
        v=vacancy
        doc.add_paragraph(f"기간 {v['t0_date']} → {v['t1_date']}", style="List Bullet")
        doc.add_paragraph(f"폐업(공실화) {v['closed']}개 · 신규 {v['opened']}개 · 순증감 {v['net']:+d}개", style="List Bullet")
        doc.add_paragraph(f"폐업률(빈 점포 발생률) {v['closure_rate_pct']}%", style="List Bullet")
    elif vacancy:
        doc.add_paragraph(f"기준 스냅샷 저장 완료: {vacancy['base_date']} · 반경 내 {vacancy['base_count']}개 업소", style="List Bullet")
        doc.add_paragraph("상가정보 분기 갱신 시 재실행하면 폐업·신규·순증감·폐업률(공실 대리지표)이 자동 산출됩니다.", style="List Bullet")
    if vacancy and vacancy.get("gonga") is not None:
        doc.add_paragraph(f"공가율(빈집) 등급: {vacancy['gonga']}/10  ※주거 빈집 지표(상업 공실과 구분)", style="List Bullet")
    doc.add_paragraph()
    # 3. 쇠퇴 등급
    _h(doc,f"3. 쇠퇴진단 등급 ({grades['area']} · {grades['year']}년 · 실데이터)")
    note=doc.add_paragraph().add_run("※ 등급값(1~10)은 전국 상대 위치. 지표별 방향(높을수록 양호/쇠퇴)이 달라 종합 판정은 공식 참고문서 범례 확정 후 산출. 아래는 원자료.")
    note.font.size=Pt(8.5); note.italic=True; note.font.color.rgb=RED
    for sector in ["인구사회","산업경제","물리환경"]:
        rows=[g for g in grades["items"] if g["sector"]==sector]
        if not rows: continue
        sp=doc.add_paragraph().add_run(f"[{sector}부문] {len(rows)}개 지표")
        sp.bold=True; sp.font.color.rgb=NAVY; sp.font.size=Pt(11)
        tb=doc.add_table(rows=1,cols=2); tb.style="Light List Accent 1"
        tb.rows[0].cells[0].text="지표"; tb.rows[0].cells[1].text="등급(1~10)"
        for g in rows:
            c=tb.add_row().cells
            c[0].text=g["mean"].replace(" 등급",""); c[1].text=str(int(g["value"]) if g["value"] else "-")
        doc.add_paragraph()
    doc.save(path); return path

def _h(doc,text):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY
